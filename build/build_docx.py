# -*- coding: utf-8 -*-
"""
build_docx.py — build a Word (.docx) copy of the CV that mirrors the LaTeX PDF
layout (large name, two-column contact row, ruled section headings, bold-lead
bullets, right-aligned dates).

DATA-DRIVEN: all content now comes from data/resume.json (previously hardcoded).
The layout helpers are unchanged from the hand-tuned original.
"""
import json
import os

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
DATA = os.path.join(ROOT, "data", "resume.json")
OUT = os.path.join(ROOT, "cv", "Nishanur_Rahman_CV.docx")

MONTHS = {
    "01": "Jan", "02": "Feb", "03": "Mar", "04": "Apr", "05": "May", "06": "Jun",
    "07": "Jul", "08": "Aug", "09": "Sep", "10": "Oct", "11": "Nov", "12": "Dec",
}


def fmt_date(value):
    if not value:
        return ""
    parts = str(value).split("-")
    if len(parts) >= 2 and parts[1] in MONTHS:
        return f"{MONTHS[parts[1]]} {parts[0]}"
    return parts[0]


def date_range(job):
    if job.get("x_dateDisplay"):
        return job["x_dateDisplay"]
    start = fmt_date(job.get("startDate"))
    end = fmt_date(job.get("endDate")) if job.get("endDate") else "Present"
    return f"{start} - {end}"


# ----------------------------------------------------------------------------
# Document + base style
# ----------------------------------------------------------------------------
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.0

for s in doc.sections:
    s.top_margin = Inches(0.5)
    s.bottom_margin = Inches(0.5)
    s.left_margin = Inches(0.5)
    s.right_margin = Inches(0.5)


def set_cell_margins(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        if side in kwargs:
            node = OxmlElement(f"w:{side}")
            node.set(qn("w:w"), str(kwargs[side]))
            node.set(qn("w:type"), "dxa")
            m.append(node)
    tcPr.append(m)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def add_two_col_row(left_runs, right_text, left_bold=True, left_italic=False,
                    right_italic=False, size=10.5, space_after=0):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    no_table_borders(table)
    table.columns[0].width = Inches(5.3)
    table.columns[1].width = Inches(2.2)
    lc, rc = table.rows[0].cells
    lc.width = Inches(5.3)
    rc.width = Inches(2.2)
    set_cell_margins(lc, left=0, right=0, top=0, bottom=0)
    set_cell_margins(rc, left=0, right=0, top=0, bottom=0)

    lp = lc.paragraphs[0]
    lp.paragraph_format.space_after = Pt(space_after)
    if isinstance(left_runs, str):
        left_runs = [(left_runs, left_bold, left_italic)]
    for text, b, i in left_runs:
        r = lp.add_run(text)
        r.bold = b
        r.italic = i
        r.font.size = Pt(size)

    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(space_after)
    rr = rp.add_run(right_text)
    rr.italic = right_italic
    rr.font.size = Pt(size)
    return table


def add_section(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "21")
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_bullet(title, body):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    if title:
        r = p.add_run(title + ": ")
        r.bold = True
        r.font.size = Pt(10.5)
    if body:
        r2 = p.add_run(body)
        r2.font.size = Pt(10.5)
    return p


def add_plain_italic_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10.5)
    return p


# ----------------------------------------------------------------------------
# Render from data
# ----------------------------------------------------------------------------
def build(d):
    b = d["basics"]
    profiles = {p["network"]: p for p in b["profiles"]}
    linkedin = profiles["LinkedIn"]["url"]
    github = profiles["GitHub"]["url"]
    site = b.get("url", "")

    # Heading
    add_two_col_row([(b["name"], True, False)], f"Email: {b['email']}", size=16, space_after=0)
    add_two_col_row([(f"LinkedIn: {linkedin}", False, False)],
                    f"Mobile:  {b['phone']}", left_bold=False, size=10.5, space_after=0)
    add_two_col_row([(f"Portfolio: {site}", False, False)],
                    f"GitHub: {github}", left_bold=False, size=10.5, space_after=2)

    # Summary
    add_section("Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run(b["summary"]).font.size = Pt(10.5)

    # Skills
    add_section("Skills")
    for g in d["skills"]:
        add_bullet(g["name"], ", ".join(g["keywords"]))

    # Experience
    add_section("Experience")
    for job in d["work"]:
        pos = job["position"]
        if job.get("x_employment") and job["x_employment"] not in (job.get("x_dateDisplay") or ""):
            pos = f"{pos} ({job['x_employment']})"
        add_two_col_row([(job["name"], True, False)], job.get("location", ""), space_after=0)
        add_two_col_row([(pos, False, True)], date_range(job),
                        left_bold=False, right_italic=True, space_after=2)
        for hi in job.get("highlights", []):
            add_bullet(hi["name"], hi["description"])

    # Hackathons
    if d.get("x_hackathons"):
        HACK_URL = d["x_hackathons"][0]["url"]
        add_section("Hackathons")
        for h in d["x_hackathons"]:
            hack_row = add_two_col_row([(h["name"], True, False)], "", space_after=0)
            _rp = hack_row.rows[0].cells[1].paragraphs[0]
            add_hyperlink(_rp, h["url"], h["event"])
            add_two_col_row([(h["tagline"], False, True)], h["year"],
                            left_bold=False, right_italic=True, space_after=2)
            for hi in h["highlights"]:
                add_bullet(hi["name"], hi["description"])
            _sub = add_bullet("Submission", h["status"] + " ")
            add_hyperlink(_sub, h["url"], "View hackathon idea")
            _sub.add_run(".").font.size = Pt(10.5)

    # Education
    add_section("Education")
    for e in d["education"]:
        degree = f"{e['studyType']} - {e['area']}"
        if e.get("score"):
            degree = f"{degree}; {e['score']}"
        right = e.get("x_dateDisplay") or fmt_date(e.get("endDate"))
        add_two_col_row([(e["institution"], True, False)], e.get("location", ""), space_after=0)
        add_two_col_row([(degree, False, True)], right,
                        left_bold=False, right_italic=True, space_after=2)
        if e.get("courses"):
            add_plain_italic_note("Relevant coursework: " + ", ".join(e["courses"]) + ".")

    # Certifications & Awards
    add_section("Certifications and Awards")
    cert_names = "; ".join(c["name"] for c in d["certificates"])
    add_bullet("Certifications", f"{cert_names}. (Verifiable on my LinkedIn: {linkedin})")
    for a in d["awards"]:
        when = fmt_date(a.get("date"))
        body = f"{a['awarder']} — {a['summary']} {when}." if when else f"{a['awarder']} — {a['summary']}"
        add_bullet(a["title"], body)

    # Publications
    if d.get("publications"):
        add_section("Publications")
        for pub in d["publications"]:
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.left_indent = Inches(0.25)
            r = para.add_run(pub["name"] + ": ")
            r.bold = True
            r.font.size = Pt(10.5)
            r2 = para.add_run(pub["summary"] + " ")
            r2.italic = True
            r2.font.size = Pt(10.5)
            if pub.get("url"):
                add_hyperlink(para, pub["url"], "View Publication")
                para.add_run(".").font.size = Pt(10.5)

    # Volunteer & Hobbies
    add_section("Volunteer Experience and Hobbies")
    for v in d["volunteer"]:
        add_two_col_row([(v["organization"], True, False)], v.get("location", ""), space_after=0)
        add_two_col_row([(v["summary"], False, True)], v.get("x_dateDisplay", ""),
                        left_bold=False, right_italic=True, space_after=2)
    hobbies = next((g["keywords"] for g in d["interests"] if g["name"] == "Hobbies"), [])
    if hobbies:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run("Hobbies: ")
        r.bold = True
        r.font.size = Pt(10.5)
        p.add_run(", ".join(hobbies)).font.size = Pt(10.5)


def main():
    with open(DATA, encoding="utf-8") as f:
        d = json.load(f)
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    build(d)
    doc.save(OUT)
    print(f"[build_docx] wrote {OUT}")


if __name__ == "__main__":
    main()
